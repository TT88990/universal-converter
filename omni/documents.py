from pathlib import Path

import fitz
import markdown as md_lib
from docx import Document as DocxDocument

from omni.errors import ConversionError
from omni.images import images_to_pdf


def _require_input(src: Path) -> None:
    if not src.is_file():
        raise ConversionError(f"Input file not found: {src}")


def _render_text_pdf(text: str, dst: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    lines_per_page = 40
    lines = text.splitlines() or [""]
    for start in range(0, len(lines), lines_per_page):
        if start > 0:
            page = doc.new_page()
        chunk = "\n".join(lines[start : start + lines_per_page])
        page.insert_textbox(fitz.Rect(72, 72, 540, 780), chunk, fontsize=11)
    doc.set_metadata(
        {"title": "UniversalConverter", "author": "UniversalConverter", "creator": "UniversalConverter", "producer": "UniversalConverter"}
    )
    doc.save(dst)
    doc.close()


def pdf_to_png(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        doc = fitz.open(src_path)
        try:
            if doc.page_count == 0:
                raise ConversionError("PDF has no pages")
            page = doc[0]
            pix = page.get_pixmap(dpi=150)
            pix.save(str(dst_path))
        finally:
            doc.close()
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"PDF to PNG failed: {exc}") from exc


def pdf_to_txt(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        doc = fitz.open(src_path)
        text = "\n\n".join(page.get_text() for page in doc)
        doc.close()
        dst_path.write_text(text, encoding="utf-8")
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"PDF to TXT failed: {exc}") from exc


def pdf_to_docx(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        doc = fitz.open(src_path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        out = DocxDocument()
        for para in text.splitlines():
            if para.strip():
                out.add_paragraph(para)
        out.save(dst_path)
    except Exception as exc:
        raise ConversionError(f"PDF to DOCX failed: {exc}") from exc


def pdf_to_md(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        doc = fitz.open(src_path)
        text = "\n\n".join(page.get_text() for page in doc)
        doc.close()
        dst_path.write_text(f"# Converted from PDF\n\n{text}\n", encoding="utf-8")
    except Exception as exc:
        raise ConversionError(f"PDF to MD failed: {exc}") from exc


def docx_to_txt(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        d = DocxDocument(str(src_path))
        text = "\n".join(p.text for p in d.paragraphs)
        dst_path.write_text(text, encoding="utf-8")
    except Exception as exc:
        raise ConversionError(f"DOCX to TXT failed: {exc}") from exc


def docx_to_md(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        d = DocxDocument(str(src_path))
        parts = ["# Document", ""]
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
                parts.append("")
        dst_path.write_text("\n".join(parts), encoding="utf-8")
    except Exception as exc:
        raise ConversionError(f"DOCX to MD failed: {exc}") from exc


def docx_to_pdf(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        d = DocxDocument(str(src_path))
        text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
        _render_text_pdf(text, dst_path)
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"DOCX to PDF failed: {exc}") from exc


def text_to_pdf(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        _render_text_pdf(src_path.read_text(encoding="utf-8"), dst_path)
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f"TXT to PDF failed: {exc}") from exc


def txt_to_docx(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        text = src_path.read_text(encoding="utf-8")
        out = DocxDocument()
        for para in text.splitlines():
            out.add_paragraph(para)
        out.save(dst_path)
    except Exception as exc:
        raise ConversionError(f"TXT to DOCX failed: {exc}") from exc


def _text_to_html(src_path: Path, dst_path: Path) -> None:
    text = src_path.read_text(encoding="utf-8")
    body = md_lib.markdown(text) if src_path.suffix.lower() == ".md" else text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>\n")
    dst_path.write_text(
        f'<!doctype html><html><meta charset="utf-8"><body>{body}</body></html>',
        encoding="utf-8",
    )


def md_to_html(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        _text_to_html(src_path, dst_path)
    except Exception as exc:
        raise ConversionError(f"MD to HTML failed: {exc}") from exc


def text_to_html(src, dst) -> None:
    src_path, dst_path = Path(src), Path(dst)
    _require_input(src_path)
    try:
        _text_to_html(src_path, dst_path)
    except Exception as exc:
        raise ConversionError(f"TXT to HTML failed: {exc}") from exc