from pathlib import Path

import fitz
import pytest
from docx import Document as DocxDocument

from omni.errors import ConversionError
from omni.documents import (
    docx_to_md,
    docx_to_pdf,
    docx_to_txt,
    md_to_html,
    pdf_to_docx,
    pdf_to_md,
    pdf_to_txt,
    text_to_html,
    text_to_pdf,
    txt_to_docx,
)

SAMPLE_TEXT = "Hello OmniConvert!\nSecond line with ünicode."


def make_pdf(path: Path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), SAMPLE_TEXT)
    doc.save(path)
    doc.close()
    return path


def make_docx(path: Path):
    d = DocxDocument()
    d.add_paragraph("Hello OmniConvert!")
    d.add_paragraph("Second line with ünicode.")
    d.save(path)
    return path


def test_pdf_to_txt(tmp_path):
    src = make_pdf(tmp_path / "in.pdf")
    dst = tmp_path / "out.txt"
    pdf_to_txt(src, dst)
    assert "Hello OmniConvert" in dst.read_text(encoding="utf-8")


def test_pdf_to_docx(tmp_path):
    src = make_pdf(tmp_path / "in.pdf")
    dst = tmp_path / "out.docx"
    pdf_to_docx(src, dst)
    docx_to_txt(dst, tmp_path / "round.txt")
    text = (tmp_path / "round.txt").read_text(encoding="utf-8")
    assert "Hello OmniConvert" in text


def test_docx_to_txt(tmp_path):
    make_docx(tmp_path / "in.docx")
    dst = tmp_path / "out.txt"
    docx_to_txt(tmp_path / "in.docx", dst)
    assert "Second line" in dst.read_text(encoding="utf-8")


def test_docx_to_md(tmp_path):
    make_docx(tmp_path / "in.docx")
    docx_to_md(tmp_path / "in.docx", tmp_path / "out.md")
    assert (tmp_path / "out.md").read_text(encoding="utf-8").startswith("#")


def test_docx_to_pdf_roundtrip(tmp_path):
    make_docx(tmp_path / "in.docx")
    docx_to_pdf(tmp_path / "in.docx", tmp_path / "out.pdf")
    assert (tmp_path / "out.pdf").stat().st_size > 1000


def test_txt_to_docx(tmp_path):
    src = tmp_path / "in.txt"
    src.write_text(SAMPLE_TEXT, encoding="utf-8")
    txt_to_docx(src, tmp_path / "out.docx")
    d = DocxDocument(str(tmp_path / "out.docx"))
    assert any("Hello OmniConvert" in p.text for p in d.paragraphs)


def test_text_to_pdf(tmp_path):
    src = tmp_path / "in.txt"
    src.write_text(SAMPLE_TEXT, encoding="utf-8")
    text_to_pdf(src, tmp_path / "out.pdf")
    assert (tmp_path / "out.pdf").stat().st_size > 1000


def test_md_and_text_to_html(tmp_path):
    md = tmp_path / "in.md"
    md.write_text("# Title\n\npara", encoding="utf-8")
    md_to_html(md, tmp_path / "out.html")
    html = (tmp_path / "out.html").read_text(encoding="utf-8")
    assert "<h1" in html
    text_to_html(tmp_path / "out.html", tmp_path / "out2.html")


def test_missing_source_raises(tmp_path):
    with pytest.raises(ConversionError):
        pdf_to_txt(tmp_path / "nope.pdf", tmp_path / "x.txt")